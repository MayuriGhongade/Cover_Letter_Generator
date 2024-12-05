import openai
import streamlit as st
import fitz  # PyMuPDF for PDF text extraction
from docx import Document
import io

# Set up OpenAI API key (you should have your own API key)
openai.api_key = "OPENAI_API_KEY"
# Function to extract text from a PDF resume
def extract_text_from_pdf(uploaded_file):
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# Function to extract text from a Word resume
def extract_text_from_word(uploaded_file):
    doc = Document(io.BytesIO(uploaded_file.read()))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Function to generate the cover letter using the new chat-based API
def generate_cover_letter(job_desc, resume_text):
    prompt = f"Create a cover letter for the following job description and resume:\n\nJob Description: {job_desc}\n\nResume: {resume_text}\n\nGenerate a professional and tailored cover letter."
    
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",  # Use the newer GPT-3.5-turbo model
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt},
        ],
        max_tokens=500
    )
    
    cover_letter = response['choices'][0]['message']['content'].strip()
    return cover_letter

# Function to save the cover letter as a .docx file
def save_as_docx(cover_letter):
    doc = Document()
    doc.add_heading('Cover Letter', 0)
    doc.add_paragraph(cover_letter)
    doc.save("cover_letter.docx")

# Streamlit app
def main():
    st.title("Cover Letter Generator")
    
    # Input for job description
    job_desc = st.text_area("Enter Job Description:")
    
    # Upload Resume (PDF or Word file)
    uploaded_file = st.file_uploader("Upload Resume (PDF or Word format)", type=["pdf", "docx"])
    
    # Generate cover letter on button click
    if st.button("Generate Cover Letter"):
        if uploaded_file and job_desc:
            # Extract resume text from PDF or Word file
            if uploaded_file.type == "application/pdf":
                resume_text = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                resume_text = extract_text_from_word(uploaded_file)
            
            # Generate cover letter using the extracted resume text and job description
            cover_letter = generate_cover_letter(job_desc, resume_text)
            
            st.subheader("Generated Cover Letter")
            st.write(cover_letter)
            
            # Export options
            if st.button("Download as DOCX"):
                save_as_docx(cover_letter)
                st.success("Cover letter saved as 'cover_letter.docx'")
        else:
            st.error("Please upload a resume and enter a job description.")
            
if __name__ == "__main__":
    main()
